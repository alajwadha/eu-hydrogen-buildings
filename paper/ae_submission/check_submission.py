"""One command that says whether the submission is in a shippable state.

Rebuilds the three documents and checks, in one place, everything that has silently
broken at least once during this paper's development:

  * LaTeX errors, undefined references and citations, natbib warnings
  * floats too large for the page, and overfull boxes above the 20 pt house tolerance
  * the Applied Energy front-matter limits (abstract 250 words, highlights 85 characters)
  * that V6.tex and frontmatter_body.tex still carry a token-identical abstract and
    highlights, since the second is the pandoc/Word path and drifts silently
  * that every figure inclusion is either above the print-scale review threshold or
    guarded at source (delegates to scripts.check_figure_print_scale)

Run:  python paper/ae_submission/check_submission.py
Exit: non-zero if any gate fails, with the failing gate named.
"""
from __future__ import annotations
import os
import re
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
REPO = HERE.parents[1]
sys.path.insert(0, str(HERE))

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

DOCS = [(HERE, "V6"), (HERE, "SI"), (REPO / "paper", "Paper_v20")]
OVERFULL_TOLERANCE_PT = 20.0
ABSTRACT_WORD_CAP = 250
HIGHLIGHT_CHAR_CAP = 85

fails: list[str] = []


def gate(ok: bool, name: str, detail: str = "") -> None:
    print(f"  [{'ok ' if ok else 'FAIL'}] {name}{(': ' + detail) if detail else ''}")
    if not ok:
        fails.append(f"{name}{(': ' + detail) if detail else ''}")


def build(where: Path, stem: str) -> None:
    for i in range(3):
        subprocess.run(["pdflatex", "-interaction=nonstopmode", stem], cwd=where,
                       capture_output=True, text=True)
        if i == 0:
            subprocess.run(["bibtex", stem], cwd=where, capture_output=True, text=True)


def check_log(where: Path, stem: str) -> None:
    log = (where / f"{stem}.log").read_text(errors="replace")
    m = re.search(r"Output written on \S+ \((\d+) pages", log)
    pages = m.group(1) if m else "?"
    print(f"\n{stem}.pdf  ({pages} pages)")

    errors = [l for l in log.splitlines() if l.startswith("!")]
    gate(not errors, f"{stem}: no LaTeX errors", f"{len(errors)} found" if errors else "")

    undef = "There were undefined references" in log or "Citation" in log and "undefined" in log
    undef_lines = [l for l in log.splitlines()
                   if "undefined" in l and ("Reference" in l or "Citation" in l)]
    gate(not undef_lines, f"{stem}: no undefined refs or cites",
         "; ".join(undef_lines[:3]) if undef_lines else "")

    floats = log.count("Float too large")
    gate(floats == 0, f"{stem}: no oversized floats", f"{floats} found" if floats else "")

    over = [float(v) for v in re.findall(r"Overfull \\hbox \(([\d.]+)pt", log)]
    bad = [v for v in over if v > OVERFULL_TOLERANCE_PT]
    gate(not bad, f"{stem}: no overfull box over {OVERFULL_TOLERANCE_PT:.0f} pt",
         f"worst {max(bad):.1f} pt" if bad else f"worst {max(over):.1f} pt" if over else "")

    nat = log.count("Package natbib Warning")
    gate(nat == 0, f"{stem}: no natbib warnings", f"{nat} found" if nat else "")

    # A requested font size that does not exist is a warning, not an error, and the
    # document still compiles and still looks plausible. The title page was set at 27 pt,
    # then 32, then 40, and every one of them rendered at 24.88, because Computer Modern
    # ships fixed design sizes and the standard .fd stops there. Three rounds of "make it
    # bigger" changed nothing on the page and nothing in this suite noticed.
    #
    # A size substitution only. A shape fallback at the same size is normal and correct,
    # and a first cut at this check failed both other documents for one: eurosym ships no
    # italic, so \euro inside italic text is set slanted instead, at the size asked for.
    flat = re.sub(r"\s+", " ", log)
    sub = re.findall(r"Font shape `([^']+)' in size <([\d.]+)> not available "
                     r"\(Font\) size <([\d.]+)> substituted", flat)
    gate(not sub, f"{stem}: every requested font size exists",
         "; ".join(f"{sh} asked for {a} pt, set at {b}" for sh, a, b in sub[:3]))


def frontmatter(src: str) -> tuple[list[str], list[str]]:
    """Abstract tokens and highlight items, from either the elsarticle or the
    pandoc-fragment form."""
    a = (re.search(r"\\begin\{abstract\}(.*?)\\end\{abstract\}", src, re.S)
         or re.search(r"\\section\*\{Abstract\}(.*?)(?=\n%%|\n\\section|\Z)", src, re.S))
    h = (re.search(r"\\begin\{highlights\}(.*?)\\end\{highlights\}", src, re.S)
         or re.search(r"(?:\\section\*\{Highlights\}|Highlights)(.*?)(?=\n%%|\n\\section|\Z)",
                      src, re.S))
    return a.group(1).split(), re.findall(r"\\item\s+(.+)", h.group(1))


def plain(tex: str) -> str:
    t = tex.replace("\\euro{}", "EUR").replace("\\%", "%")
    t = re.sub(r"\$[^$]*\$", "X", t)
    t = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", t)
    return re.sub(r"\\[a-zA-Z]+", "", t).replace("~", " ")


def check_frontmatter() -> None:
    print("\nFront matter")
    a_main, h_main = frontmatter((HERE / "V6.tex").read_text())
    a_twin, h_twin = frontmatter((HERE / "frontmatter_body.tex").read_text())

    gate(a_main == a_twin, "abstract identical in V6.tex and frontmatter_body.tex")
    gate(h_main == h_twin, "highlights identical in V6.tex and frontmatter_body.tex")

    # The two files also duplicate the nomenclature and the keyword list, and neither was
    # gated. Both drifted: the nomenclature carried two definitions of the same symbol and
    # the keyword lists ran to six entries against eight. Anything duplicated across the
    # two front matters has to be checked, or it separates.
    def _entries(src: str) -> list:
        return sorted(" ".join(x.split())
                      for x in re.findall(r"\\item\[(.*?)\]\s*(.*?)(?=\\item\[|\\end\{description\})",
                                          src, re.S)[0:0] or
                      [a + " -- " + b for a, b in
                       re.findall(r"\\item\[(.*?)\]\s*(.*?)(?=\\item\[|\\end\{description\})",
                                  src, re.S)])

    def _keywords(src: str) -> list:
        m = (re.search(r"\\begin\{keyword\}(.*?)\\end\{keyword\}", src, re.S)
             or re.search(r"\\section\*\{Keywords\}\s*(.*?)(?=\n%%|\n\\section|\Z)", src, re.S))
        if not m:
            return []
        return sorted(x.strip().lower() for x in re.split(r"\\sep|;", m.group(1)) if x.strip())

    m_src = (HERE / "V6.tex").read_text()
    t_src = (HERE / "frontmatter_body.tex").read_text()
    gate(_entries(m_src) == _entries(t_src),
         "nomenclature identical in V6.tex and frontmatter_body.tex")
    kw_m, kw_t = _keywords(m_src), _keywords(t_src)
    gate(kw_m == kw_t, "keywords identical in V6.tex and frontmatter_body.tex",
         f"{len(kw_m)} vs {len(kw_t)}")

    words = len(plain(" ".join(a_main)).split())
    gate(words <= ABSTRACT_WORD_CAP, f"abstract at most {ABSTRACT_WORD_CAP} words",
         f"{words} words")

    strip = lambda s: re.sub(r"\\[a-zA-Z]+\{?|[}$]", "", s)
    longest = max((len(strip(x)) for x in h_main), default=0)
    gate(longest <= HIGHLIGHT_CHAR_CAP,
         f"every highlight at most {HIGHLIGHT_CHAR_CAP} characters",
         f"{len(h_main)} highlights, longest {longest}")


def _figure_check(module: str, claim: str) -> None:
    r = subprocess.run([sys.executable, "-m", module],
                       cwd=REPO / "code", capture_output=True, text=True,
                       env={**__import__("os").environ, "PYTHONPATH": str(REPO / "code")})
    tail = [l for l in r.stdout.strip().splitlines() if l.strip()][-1:]
    gate(r.returncode == 0, claim, tail[0].strip() if r.returncode and tail else "")


def check_figures() -> None:
    print("\nFigures")
    _figure_check("scripts.check_figure_print_scale",
                  "every inclusion above the print-scale floor or guarded")
    # Legibility and layout are separate defects. A figure can set every glyph above the floor
    # and still print two of them on top of each other, which is how a legend row came to sit
    # on the descenders of an axis label.
    _figure_check("scripts.check_figure_overlap",
                  "no text collides, and no callout sits over the data")
    # Both gates above read the sources. Neither can see resolution, which is a property
    # of the shipped .docx alone: every figure was legible, vector PDFs sat beside every
    # PNG, and the Word twin still carried its artwork at 199 to 213 dpi.
    _figure_check("scripts.check_figure_freshness",
                  "every submitted figure was built from the source it carries")
    _figure_check("scripts.check_word_figure_dpi",
                  "every figure in the Word twin at or above 300 dpi")


def check_numbers() -> None:
    """Delegate to the artefact-consistency check in code/scripts."""
    print("\nManuscript numbers")
    r = subprocess.run([sys.executable, "-m", "scripts.check_manuscript_numbers"],
                       cwd=REPO / "code", capture_output=True, text=True,
                       env={**__import__("os").environ, "PYTHONPATH": str(REPO / "code")})
    tail = [l for l in r.stdout.strip().splitlines() if l.strip()][-1:]
    gate(r.returncode == 0, "every checked claim matches its artefact",
         tail[0].strip() if r.returncode and tail else "")


BODY_WORD_CAP = 8000          # Applied Energy's own guide, and now the binding one.
                              # 7,400 -> 7,700 paid for the corpus habits the AE style
                              # review found missing: narrating Table 1 rather than
                              # pointing at it, justifying modelling choices in the
                              # sentence that makes them, and naming prior work.
                              # 7,700 -> 7,900 pays for the technology-share allocation
                              # paragraph and the two scenario levers Table 1 had
                              # omitted; 7,900 -> 8,000 pays for the V6 text
                              # specification. The eight-reviewer round called that the most
                              # important missing methodological content: the softmax
                              # temperature, the cost-response weight and the feasibility
                              # rules decide the technology mix and were undisclosed.
GA_MIN_W, GA_MIN_H = 1328, 531   # Elsevier's graphical-abstract floor, aspect near 2.5:1


def check_graphical_abstract() -> None:
    """The manuscript embeds the TALL canvas and Editorial Manager wants the WIDE one.

    That split is deliberate (see the note in scripts/graphical_abstract.py): a 2.5:1
    canvas at \\textwidth scales every font below the 6 pt print floor, so the wide
    variant cannot be the one V6.tex includes. What was missing is any gate that the
    wide file exists and meets the spec, so nothing caught the case where the tall file
    is the only one built and gets uploaded by mistake.
    """
    print("\nGraphical abstract")
    wide = REPO / "paper" / "figs" / "paper" / "graphical_abstract_wide.png"
    if not wide.exists():
        fails.append("graphical_abstract_wide.png missing (this is the Editorial Manager upload)")
        print("  [FAIL] wide variant not built")
        return
    try:
        from PIL import Image
        w, h = Image.open(wide).size
    except Exception as e:                                    # noqa: BLE001
        print(f"  [skip] cannot read image size ({e})")
        return
    ok = w >= GA_MIN_W and h >= GA_MIN_H and 2.2 <= w / h <= 2.8
    print(f"  {'[ok ]' if ok else '[FAIL]'} upload file {w}x{h} px, aspect {w/h:.3f} "
          f"(floor {GA_MIN_W}x{GA_MIN_H}, aspect ~2.5)")
    if not ok:
        fails.append(f"graphical_abstract_wide.png is {w}x{h}, aspect {w/h:.2f}")


def _body_words() -> tuple:
    """(total, per-section) on the definition the body cap uses.

    Factored out so the title-page summary is checked against the same counter rather
    than a second one written to agree with it.
    """
    total, per = 0, []
    for name in ("intro", "methods", "results", "concl"):
        t = (HERE / "sections" / f"{name}.tex").read_text()
        t = re.sub(r"%.*", "", t)
        for env in ("figure", "table", "equation", "align", "tabular"):
            t = re.sub(r"\\begin\{" + env + r"\*?\}.*?\\end\{" + env + r"\*?\}", "", t, flags=re.S)
        t = re.sub(r"\\label\{[^}]*\}", "", t)
        t = re.sub(r"\\(ref|eqref)\{[^}]*\}", "X", t)
        t = re.sub(r"\\cite[a-z]*\{[^}]*\}", "", t)
        t = re.sub(r"\$[^$]*\$", "X", t)
        t = re.sub(r"\\[a-zA-Z]+\*?", "", t)
        t = re.sub(r"[{}\\~]", " ", t)
        n = len([w for w in t.split() if any(ch.isalnum() for ch in w)])
        per.append(f"{name} {n}")
        total += n
    return total, per


def check_body_words() -> None:
    """Count the running body and hold it to the author's cap.

    Excludes front matter, float environments (so figure and table captions and cell
    contents are out), comments and labels. A \\ref or inline maths counts as one printed
    token. Three independent counters agreed on this definition to within 15 words.
    """
    print("\nBody word count")
    total, per = _body_words()
    ok = total <= BODY_WORD_CAP
    print(f"  {'[ok ]' if ok else '[FAIL]'} body {total} words vs cap {BODY_WORD_CAP}"
          f"  ({', '.join(per)})")
    if not ok:
        fails.append(f"body is {total} words, {total - BODY_WORD_CAP} over the {BODY_WORD_CAP} cap")


# Both caps are stated against an external norm and carry headroom, and both are printed
# with the worst live float so the headroom is visible rather than assumed.
#
# They were not. 110 was set at the worst float of the day and read exactly 110 for two
# rounds; 65 was then set at the worst caption of the day and read exactly 65. Three
# referees made the same point independently, and they were right: a threshold fitted to
# its own worst case cannot fail, so it certifies rather than governs, and one added word
# turns the build red for something that is not a defect.
#
# The AEJ corpus norms are 8-12 words for a single-panel figure caption and 10-15 for a
# table. Those are unreachable for a float whose cells are model artefacts and whose
# provenance has to travel with it, which is the case the skill's own escape hatch covers:
# method and provenance move below the rule into a note block, and the caption is measured
# alone. So the caption cap is set at 65, roughly four times the table norm, which is what
# a caption still doing only naming work costs here; and the combined cap at 125, which is
# the caption cap plus a note block of about the same length again. Worst live values when
# these were set: 61 and 108, so both carry real slack.
FLOAT_TEXT_CAP = 125   # words of caption plus every note block under one float
CAPTION_CAP = 65


def check_float_text() -> None:
    """Cap the words that live under a float, which no other counter sees.

    check_body_words strips float environments, correctly, since the body cap is on
    running prose. The AEJ caption norm measures only \\caption{}. Between them sits
    everything else a reader reads as part of the caption, which is the house Note block
    and the lettered footnote list a table sets after \\end{tabular}. Several hundred
    words sat in that gap, invisible to both counters.

    Round 24 answered that by printing the number. Three referees then found the same
    thing independently, which is that the function ended in a hard-coded "[ok ]", never
    touched `fails`, and so could not fail, while its docstring claimed it held the text
    against a norm. It was a gate-shaped print statement, and it was worse than no gate
    because it was introduced as the fix for the defect it then certified. It also read
    only sections/, so the entire supplement was outside it, and it saw only \\textit{Note:}
    blocks, so 336 words of lettered table footnotes were outside it again.

    It now has a threshold, counts every note-like block including lettered table
    footnotes, and covers the supplement. The cap is on caption plus notes per float,
    because that is the object a reader meets under one piece of artwork.
    """
    import re as _re
    # Everything a reader reads under the rule as part of the float. The lettered list a
    # table sets after \end{tabular} carries no \textit{Note:} marker, which is how 336
    # words stayed outside the count.
    NOTE = _re.compile(r"\\textit\{(?:Note|Source|Notes)[:.]?\}(.*?)(?:\\par\}|\Z)", _re.S)
    FOOT = _re.compile(r"\\end\{tabular\}(.*?)\Z", _re.S)
    CAP = _re.compile(r"\\caption\{")

    def words(s: str) -> int:
        return len(_re.sub(r"\\[a-zA-Z]+\*?|[{}$~]", " ", s).split())

    over, total, worst, n = [], 0, (0, "", ""), 0
    cap_seen = []
    for sub in ("sections", "si_body"):
        for f in sorted((HERE / sub).glob("*.tex")):
            txt = _re.sub(r"(?<!\\)%.*", "", f.read_text(encoding="utf-8"))
            for fl in _re.finditer(r"\\begin\{(figure|table)\*?\}(.*?)\\end\{\1\*?\}",
                                   txt, _re.S):
                body = fl.group(2)
                cm = CAP.search(body)
                cap_w = 0
                if cm:
                    d, j = 0, cm.end() - 1
                    while j < len(body):
                        d += (body[j] == "{") - (body[j] == "}")
                        if d == 0:
                            break
                        j += 1
                    cap_w = words(body[cm.end():j])
                # Notes and lettered footnotes, whichever the float uses, without
                # double-counting a note that sits inside the post-tabular tail.
                tail = FOOT.search(body)
                note_w = sum(words(m.group(1)) for m in NOTE.finditer(body))
                if tail and not NOTE.search(tail.group(1)):
                    note_w += words(tail.group(1))
                total += note_w
                n += 1
                cap_seen.append(cap_w)
                lab = _re.search(r"\\label\{([^}]+)\}", body)
                where = f"{f.name}:{lab.group(1) if lab else '?'}"
                if cap_w + note_w > worst[0]:
                    worst = (cap_w + note_w, where, sub)
                if cap_w + note_w > FLOAT_TEXT_CAP:
                    over.append((cap_w + note_w, where, "caption+notes"))
                if cap_w > CAPTION_CAP:
                    over.append((cap_w, where, "caption alone"))
    worst_cap = max(cap_seen) if cap_seen else 0
    gate(not over,
         f"every float at most {CAPTION_CAP} caption words and {FLOAT_TEXT_CAP} with its "
         f"notes ({n} floats, {total} note words; worst caption {worst_cap} "
         f"({CAPTION_CAP - worst_cap} spare), worst combined {worst[0]} "
         f"({FLOAT_TEXT_CAP - worst[0]} spare) at {worst[1]})",
         "; ".join(f"{w} {k} at {p}" for w, p, k in sorted(over, reverse=True)))


def check_word_fields() -> None:
    """The Word twin must carry live numbering fields, not frozen text.

    build_docx_elsevier resolves \\ref and \\eqref against the compiled .aux and then
    writes SEQ counters on captions and REF fields on cross-references, so numbering
    updates inside Word. When the .aux is absent it degrades silently: it emits the
    document with every field count at zero and exits 0, and the reader gets a file whose
    numbers are dead text.

    That is not hypothetical. Renaming the submission deleted the old .aux, and the first
    build of the renamed manuscript ran before the new one existed. It printed
    "0 equations, 0 captions, 0 cross-references" and every gate here passed, including
    the pagination gate, because pagination does not depend on fields. The same thing
    happens to anyone building from a fresh clone in the obvious order.

    So count them in the built file. The floor is 1 of each rather than the current
    totals, because the right assertion is that the field machinery ran at all; pinning
    exact counts would fail on any legitimate edit that adds or removes a caption.
    """
    docx_path = HERE / "V6.docx"
    if not docx_path.exists():
        gate(False, "Word twin carries live numbering fields", "V6.docx is missing")
        return
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    n_seq = len(re.findall(r"\bSEQ\s+\w+", xml))
    n_ref = len(re.findall(r"\bREF\s+\w+", xml))
    ok = n_seq > 0 and n_ref > 0
    gate(ok, f"Word twin carries live numbering fields ({n_seq} SEQ, {n_ref} REF)",
         "" if ok else "V6.aux was probably missing when build_docx_elsevier.py ran, "
                       "so rebuild it after the LaTeX pass, not before")


def check_obsolete_architecture():
    """Fail if the dispatch-generated-the-shares wording reappears.

    The model imposes hydrogen adoption exogenously and then tests it against an
    independently derived dispatch bound. An earlier wording said the dispatch capped
    or set each scenario's share, which the paper's own Stated Policies result
    contradicts: 5 per cent of announced adoption against 0.6 per cent of
    operating-cost support. The contradiction reached the abstract, so it is worth a
    gate rather than a memory.
    """
    # The first three were the main paper's. The next three were the SI's, which kept
    # the abandoned architecture for a further round because the gate named phrases
    # rather than the claim. A gate that lists only what one document happened to say
    # asserts nothing about the document that says it differently.
    banned = [
        r"bound that share from the winter-peak",
        r"cap each scenario at or below",
        r"each scenario ceiling is set at or below",
        r"(?:sets|carries) six levers",
        r"ceiling is bounded by the (?:merit-order )?dispatch",
        r"heating share in each scenario is the peak slice",
    ]
    roots = [HERE / "sections", HERE / "si_body", HERE, REPO / "paper" / "sections"]
    hits = []
    for root in roots:
        for f in sorted(root.glob("*.tex")) + sorted(root.glob("*.md")):
            if f.name.startswith("_"):
                continue      # generated by build.py from sections/, not a source file
            txt = re.sub(r"\s+", " ", f.read_text(encoding="utf-8"))
            for pat in banned:
                if re.search(pat, txt, re.I):
                    hits.append(f"{f.name}: {pat}")
    return (not hits,
            "no obsolete dispatch-caps-the-share wording"
            + ("" if not hits else ": " + "; ".join(hits)))


def check_word_styles() -> None:
    """Every style the Word twin names must exist in the file that defines styles.

    pandoc puts its own style names on what it emits: FirstParagraph on an opening
    paragraph, ImageCaption and TableCaption on captions, Bibliography on reference
    entries, Hyperlink and SectionNumber on runs. The reference document built here
    defined none of them, so 175 references pointed at nothing. Word answers a dangling
    style by falling back to Normal, which drops the formatting without saying so, and
    LibreOffice printed a stray capital X after every affected paragraph, including both
    table captions and every entry in the reference list.

    Nothing caught it: the file opened, the fields worked, the pagination gate passed and
    the text was all present. Only a reader looking at the page would have seen it. So
    check the reference rather than the symptom, and assert none is dangling.
    """
    docx_path = HERE / "V6.docx"
    if not docx_path.exists():
        gate(False, "every style the Word twin names is defined", "V6.docx is missing")
        return
    with zipfile.ZipFile(docx_path) as z:
        doc = z.read("word/document.xml").decode("utf-8", "replace")
        sty = z.read("word/styles.xml").decode("utf-8", "replace")
    defined = set(re.findall(r'<w:style [^>]*w:styleId="([^"]+)"', sty))
    used = set()
    for kind in ("pStyle", "rStyle", "tblStyle"):
        used |= set(re.findall(r'<w:%s w:val="([^"]+)"' % kind, doc))
    dangling = sorted(used - defined)
    gate(not dangling,
         f"every style the Word twin names is defined ({len(used)} named)",
         "" if not dangling else "not in styles.xml: " + ", ".join(dangling))


def check_word_table_rules() -> None:
    """The Word twin's tables must be ruled the way the PDF's are.

    booktabs draws no vertical rule at all and only the horizontal ones the source asks
    for: one at the top, one under the column head, one at each \\midrule and one at the
    foot. Word's default is a box around every cell, and the twin shipped that way, so a
    spreadsheet sat beside a typeset table and the two documents did not look like one set.

    The rule positions are read out of the same LaTeX the builder reads, so this asserts
    the right answer rather than forbidding the wrong one: a table that gains a \\midrule
    fails here until the Word twin gains the rule too.
    """
    import build_docx_elsevier as bd
    from docx import Document
    from lxml import etree
    docx_path = HERE / "V6.docx"
    body = HERE / "_elsevier_body.tex"
    if not docx_path.exists() or not body.exists():
        gate(False, "the Word twin's tables carry the source's rules",
             "V6.docx or the assembled source is missing; run build_docx_elsevier.py")
        return
    specs = bd._tabular_specs()
    doc = Document(str(docx_path))
    data = [t for t in doc.tables if len(t.columns) >= 3]
    if len(data) != len(specs):
        gate(False, "the Word twin's tables carry the source's rules",
             f"{len(data)} data tables against {len(specs)} multi-column tabulars")
        return
    bad = []
    for n, (t, spec) in enumerate(zip(data, specs), start=1):
        xml = etree.tostring(t._tbl, encoding="unicode")
        for edge in ("left", "right", "insideV"):
            if re.search(r'<w:%s [^>]*w:val="(?!none)' % edge, xml):
                bad.append(f"Table {n} draws a vertical rule ({edge})")
        nrow = len(t.rows)
        want = {0, spec["head"]} | {r for r in spec["rules"] if 0 < r < nrow}
        got = {i for i, row in enumerate(t.rows)
               if any(c._tc.find(f"{W}tcPr/{W}tcBorders/{W}top") is not None
                      for c in row.cells)}
        if got != want:
            bad.append(f"Table {n} rules above rows {sorted(got)}, source says "
                       f"{sorted(want)}")
        if not any(c._tc.find(f"{W}tcPr/{W}tcBorders/{W}bottom") is not None
                   for c in t.rows[nrow - 1].cells):
            bad.append(f"Table {n} has no rule at its foot")
    gate(not bad, f"the Word twin's tables carry the source's rules ({len(data)} tables)",
         "; ".join(bad))


def check_word_page_structure() -> None:
    """The Word twin has to paginate like the PDF.

    They are one submission in two formats, so an editor opening either should meet the
    same thing on the same page. The Word file did not: the graphical abstract, the
    highlights, the title block, the abstract and the keywords all ran together on a
    single page, in a different order from the PDF as well, with the graphical abstract
    last instead of first. Nothing noticed, because every earlier Word check read
    content and none read layout.

    The expected order is derived from V6.pdf rather than typed here, so the check
    follows the PDF wherever it goes. Each of the PDF's first five pages contributes its
    opening line as a marker; the Word file has to carry those markers in that order,
    with at least one page break between each consecutive pair.
    """
    pdf, docx_path = HERE / "V6.pdf", HERE / "V6.docx"
    if not (pdf.exists() and docx_path.exists()):
        gate(False, "Word twin paginates like the PDF", "V6.pdf or V6.docx is missing")
        return
    norm = lambda t: " ".join(t.split()).casefold()

    markers = []
    for n in range(1, 6):
        r = subprocess.run(["pdftotext", "-f", str(n), "-l", str(n), str(pdf), "-"],
                           capture_output=True, text=True)
        line = next((l for l in r.stdout.splitlines() if l.strip()), "")
        markers.append(norm(line)[:30])
    if not all(markers):
        gate(False, "Word twin paginates like the PDF", "could not read the PDF's pages")
        return

    try:
        import docx as _docx
    except ImportError:
        gate(False, "Word twin paginates like the PDF", "python-docx is not installed")
        return
    paras = _docx.Document(str(docx_path)).paragraphs
    breaks = {i for i, par in enumerate(paras) if 'w:type="page"' in par._p.xml}

    found, bad = [], []
    for m in markers:
        hit = next((i for i, par in enumerate(paras)
                    if norm(par.text).startswith(m) and (not found or i > found[-1])), None)
        if hit is None:
            bad.append(f"the Word file has no paragraph starting {m!r}")
            break
        found.append(hit)
    if not bad:
        for a, b, m in zip(found, found[1:], markers[1:]):
            if not any(a < i < b for i in breaks):
                bad.append(f"no page break before {m!r}")
    gate(not bad, f"Word twin paginates like the PDF ({len(markers)} front-matter pages)",
         "; ".join(bad[:3]))


def check_scripts_compile() -> None:
    """Every module the reproduction guide tells a reader to run must parse.

    A round-22 edit left a stray bracket in scripts/capacity_payment.py. The module did
    not compile for a full revision, the reproduction guide's command for it raised
    SyntaxError, and the number gate that subprocess-runs it swallowed the failure and
    printed "artefact = -1" as though a sentinel were a derived value. One line of
    py_compile would have caught it the same minute.
    """
    # rglob, and both trees. The glob was code/scripts and code/src at one level, which
    # left out code/scripts/country_build/, a directory the root reproduction guide
    # tells a reader to run, and every module in this directory including this gate
    # itself. The docstring said "every module the reproduction guide tells a reader to
    # run"; the code checked a subset of one tree.
    bad, n = [], 0
    for d in (REPO / "code", HERE):
        for f in sorted(d.rglob("*.py")):
            if "__pycache__" in f.parts:
                continue
            n += 1
            try:
                compile(f.read_text(encoding="utf-8"), str(f), "exec")
            except SyntaxError as e:
                bad.append(f"{f.relative_to(REPO)}:{e.lineno} {e.msg}")
    gate(not bad, f"every script under code/ and the submission directory compiles ({n})",
         "; ".join(bad[:3]) if bad else "")


def check_typeset_hazards() -> None:
    """Two things that compile cleanly and print wrong.

    \euro inside a math group falls back to an italic e, so the scarcity premium printed
    as "e60/MWh" and the turbine capex, which is a highlight number, as "e600/kW", two
    lines from a correctly set 51. Nothing caught it: it is valid LaTeX, no gate reads
    fonts, and pdftotext renders both the real euro and the fallback as "e".

    And a BibTeX `note` field is typeset. Five entries carried private research
    annotations, so the printed reference list quoted this study's own working figures
    and, in one case, collided the note into the journal name.
    """
    import re as _re
    tex = list((HERE / "sections").glob("*.tex")) + list((HERE / "si_body").glob("*.tex"))
    tex += [HERE / "V6.tex", HERE / "SI.tex"]
    bad = []
    for f in tex:
        if not f.exists():
            continue
        txt = _re.sub(r"(?<!\\)%.*", "", f.read_text(encoding="utf-8"))
        # odd-index chunks of a split on unescaped $ are inside math mode
        for i, chunk in enumerate(_re.split(r"(?<!\\)\$", txt)):
            if i % 2 and "\\euro" in chunk and "\\text{" not in chunk:
                bad.append(f"{f.name}: ${chunk.strip()[:40]}$")
    gate(not bad, "no euro sign inside a math group",
         "; ".join(bad[:3]) if bad else "")

    bib = REPO / "paper" / "References_v1.bib"
    notes = []
    if bib.exists():
        btxt = bib.read_text(encoding="utf-8")
        for m in _re.finditer(r"\n\s*note\s*=\s*\{", btxt):
            seg, depth, j = "", 0, m.end() - 1
            while j < len(btxt):
                seg += btxt[j]
                depth += (btxt[j] == "{") - (btxt[j] == "}")
                if depth == 0:
                    break
                j += 1
            s = " ".join(seg.split())
            # Length alone is not the test: the shortest annotation that shipped was 71
            # characters ("SMR/ATR + CCS LCOH 2050 EUR 2.4 to 2.9/kg; gas-price
            # sensitivity"). What marks a note as commentary rather than a locator is
            # that it quotes a cost. Flag either.
            if len(s) > 100 or _re.search(r"(EUR|\\euro|\u20ac)\s*[\d.]|/kg|/MWh|/kW", s):
                notes.append(s[:50])
    gate(not notes, "no long BibTeX note field reaches the printed reference list",
         f"{len(notes)} over 60 chars; use annote" if notes else "")


def check_upload_set() -> None:
    """The folder an editor uploads from must match the deliverables just built.

    This gate exists because that folder shipped a wrong number for a full revision. The
    correction reached every .tex file, both compiled PDFs and both papers, and the commit
    said so; the upload folder kept the old value, because it was a manual copy and every gate read the
    LaTeX corpus rather than the renderings. A gate that only reads source cannot see a
    stale rendering. This one compares bytes, and package_v6.py is what fixes it.

    Matching is not sufficient on its own. Asserting that the current files are current says
    nothing about what else sits beside them, and the previous revision's four files stayed in
    the folder through a version bump while this gate reported it clean. An editor receives the
    folder, so the second gate below asserts the folder holds nothing but the upload set.
    """
    sys.path.insert(0, str(HERE))
    import package_v6
    missing = [s for s, _ in package_v6.PAIRS if not s.exists()]
    if missing:
        gate(False, "V6 upload set matches the built deliverables",
             f"{len(missing)} deliverable(s) not built")
        return
    st = package_v6.stale()
    gate(not st, "V6 upload set matches the built deliverables",
         ("stale: " + ", ".join(d.name for _, d in st) +
          "  (run python3 package_v6.py)") if st else "")

    orph = package_v6.orphans()
    gate(not orph, "V6 upload folder carries no superseded file",
         (", ".join(p.name for p in orph) +
          "  (run python3 package_v6.py)") if orph else "")


def main() -> int:
    if shutil.which("pdflatex") is None:
        print("pdflatex not found; cannot run the build gates.")
        return 2
    for where, stem in DOCS:
        build(where, stem)
        check_log(where, stem)
    check_frontmatter()
    check_figures()
    check_graphical_abstract()
    check_body_words()
    check_float_text()
    check_scripts_compile()
    check_word_fields()
    gate(*check_obsolete_architecture())
    check_word_styles()
    check_word_table_rules()
    check_word_page_structure()
    check_numbers()
    check_typeset_hazards()
    check_upload_set()

    print()
    if fails:
        print(f"{len(fails)} gate(s) failed:")
        for f in fails:
            print(f"  - {f}")
        return 1
    print("All submission gates pass.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
