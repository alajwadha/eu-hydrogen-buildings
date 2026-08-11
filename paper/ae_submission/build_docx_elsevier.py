#!/usr/bin/env python3
"""Build the Applied Energy submission as an Elsevier-conformant Word manuscript.

This is NOT build.py. That script makes a readable twin for circulation: A4, single
spaced, pandoc's default look. It is fine for sending a co-author something legible and
it is wrong for a journal desk.

Elsevier's review layout, which is what elsarticle's [review] option produces and what a
Word submission has to imitate, is specific:

  single column, no two-column typesetting
  Times New Roman 12 pt throughout, including captions and the reference list
  double line spacing in the body
  continuous line numbers down the left margin, restarting nowhere
  page numbers
  a title page carrying the title, the authors with affiliation superscripts, the
    affiliations, and the corresponding author with an email
  Highlights and Keywords as their own labelled blocks
  numbered section headings, figure captions BELOW figures, table captions ABOVE tables
  a numbered reference list, since Applied Energy uses the numeric style

Every one of those is set here rather than left to pandoc's defaults. The content comes
from the same gated LaTeX the submission PDF is built from, so the two cannot diverge:
frontmatter_body.tex plus sections/{intro,methods,results,concl}.tex, with citations
resolved through the same .bib and the same numeric CSL.

Run:  python3 paper/ae_submission/build_docx_elsevier.py
Out:  paper/ae_submission/V6.docx
"""
from __future__ import annotations

import os
import re
import subprocess
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
BIB = os.path.abspath(os.path.join(HERE, "..", "References_v1.bib"))
CSL = os.path.abspath(os.path.join(HERE, "..", "ieee.csl"))
FIGDIR = os.path.abspath(os.path.join(HERE, "..", "figs", "paper"))
# Named for its source, so V6.tex, V6.pdf and V6.docx are one set.
OUT = os.path.join(HERE, "V6.docx")
REF = os.path.join(HERE, "_ref_elsevier.docx")

BODY_FONT = "Times New Roman"
BODY_PT = 12
LINE_SPACING = 2.0


# ── front matter, read from V6.tex so it cannot drift from the submission ──────────
def frontmatter() -> dict:
    t = open(os.path.join(HERE, "V6.tex"), encoding="utf-8").read()

    def grab(env):
        m = re.search(r"\\begin\{%s\}(.*?)\\end\{%s\}" % (env, env), t, re.S)
        return m.group(1).strip() if m else ""

    title = re.search(r"\\title\{(.+?)\}\s*\n", t, re.S)
    # \author[tag]{Name\corref{cor1}} -- strip the corresponding-author marker,
    # which otherwise renders as "Ali Alajwadcor1" on the title page.
    # The author argument nests: \author[cornell]{Ali Alajwad\corref{cor1}}. A [^}]*
    # capture stops at the FIRST closing brace and returns "Ali Alajwad\corref{cor1",
    # which then cleans to "Ali Alajwadcor1" on the title page. Allow one nesting level.
    authors = [(tag, re.sub(r"\\corref\{[^}]*\}", "", name))
               for tag, name in re.findall(
                   r"\\author\[([^\]]*)\]\{((?:[^{}]|\{[^{}]*\})*)\}", t)]
    affils = re.findall(r"\\affiliation\[([^\]]*)\]\{organization=\{([^}]*)\}(.*?)\}\s*\n", t)
    highlights = [h.strip() for h in re.findall(r"\\item\s+(.+)", grab("highlights"))]
    keywords = [k.strip() for k in re.split(r"\\sep", grab("keyword")) if k.strip()]
    # Read the corresponding author off \corref rather than hard-coding a name and
    # address here, so changing either in V6.tex reaches the Word file.
    cor = re.search(r"\\author\[[^\]]*\]\{((?:[^{}]|\{[^{}]*\})*?)\\corref", t)
    ead = re.search(r"\\ead\{([^}]*)\}", t)
    return dict(
        title=clean(title.group(1)) if title else "",
        corresponding=clean(cor.group(1)) if cor else "",
        email=ead.group(1).strip() if ead else "",
        authors=[(tag, clean(name)) for tag, name in authors],
        affils=[(tag, _affil(org, extra)) for tag, org, extra in affils],
        abstract=clean(grab("abstract")),
        highlights=[clean(h) for h in highlights],
        keywords=[clean(k) for k in keywords],
    )


def _affil(org: str, extra: str) -> str:
    """organization={X}, city={Y}, state={Z}, country={W} -> "X, Y, Z, W".

    The field separators are already commas, so substituting ", " for each key name
    doubled every one and produced "Cornell University, , Ithaca, , NY, , USA".
    """
    bits = [clean(org)]
    for key in ("city", "state", "country"):
        m = re.search(r"%s=\{([^}]*)\}" % key, extra)
        if m and m.group(1).strip():
            bits.append(clean(m.group(1)))
    return ", ".join(bits)


def clean(s: str) -> str:
    """LaTeX to plain text for the front-matter fields only."""
    s = s.replace(r"\euro{}", "€").replace(r"\euro", "€")
    s = s.replace(r"\%", "%").replace("~", " ")
    s = re.sub(r"\\ce\{CO2\}", "CO₂", s)
    s = re.sub(r"\\text\w*\{([^}]*)\}", r"\1", s)
    s = re.sub(r"\$([^$]*)\$", r"\1", s)
    s = s.replace("_2", "₂").replace("{,}", ",")
    s = re.sub(r"\\[a-zA-Z]+\*?", "", s)
    s = s.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", s).strip()


# ── the style template pandoc renders onto ───────────────────────────────────────────
def make_reference_docx() -> None:
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    d = docx.Document()

    sec = d.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)          # A4 portrait
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Cm(2.5))

    normal = d.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_PT)
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # East-Asian and complex-script fonts have to be named too or Word substitutes.
    from docx.oxml.ns import qn
    rpr = normal.element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rpr.set(qn(attr), BODY_FONT)

    for name, size, bold in (("Heading 1", 12, True), ("Heading 2", 12, True),
                             ("Heading 3", 12, True), ("Heading 4", 12, True)):
        st = d.styles[name]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = bold
        st.font.color.rgb = None
        st.paragraph_format.line_spacing = LINE_SPACING
        st.paragraph_format.space_before = Pt(12)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rf = st.element.get_or_add_rPr().get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rf.set(qn(attr), BODY_FONT)

    # Captions: same face and size, italic off, so a reviewer's ruler finds 12 pt.
    if "Caption" in [s.name for s in d.styles]:
        cap = d.styles["Caption"]
        cap.font.name = BODY_FONT
        cap.font.size = Pt(BODY_PT)
        cap.font.italic = False
        cap.font.color.rgb = None
        cap.paragraph_format.line_spacing = LINE_SPACING

    add_pandoc_styles(d)
    d.save(REF)


# pandoc names these on the paragraphs and runs it emits, and the reference document built
# above defined none of them, so 175 style references in the converted file pointed at
# nothing. Word falls back to Normal and drops the formatting; LibreOffice prints a stray
# capital X after every affected paragraph, which is the mark that appeared after both
# table captions, all seven figure captions and every reference entry. Defining the styles
# is the fix. Deleting the pStyle attributes would hide the mark and lose the formatting
# with it.
#
# The last field says whether the style must cancel the bold it inherits. The Caption
# style is bold throughout, and a caption here is bold only in its lead-in sentence, which
# pandoc marks on the run. Inheriting the style's bold set the whole caption in bold.
PANDOC_STYLES = (
    ("FirstParagraph",  "First Paragraph",  "paragraph", "BodyText", False),
    ("Bibliography",    "Bibliography",     "paragraph", "Normal",   False),
    ("CaptionedFigure", "Captioned Figure", "paragraph", "Normal",   False),
    ("ImageCaption",    "Image Caption",    "paragraph", "Caption",  True),
    ("TableCaption",    "Table Caption",    "paragraph", "Caption",  True),
    ("SectionNumber",   "Section Number",   "character", "DefaultParagraphFont", False),
    ("Hyperlink",       "Hyperlink",        "character", "DefaultParagraphFont", False),
)


def add_pandoc_styles(d) -> int:
    """Define every style pandoc names, so no reference in the output is dangling."""
    from docx.oxml.ns import qn
    styles = d.styles.element
    have = {s.get(qn("w:styleId")) for s in styles.findall(qn("w:style"))}
    for sid, _, _, base, _ in PANDOC_STYLES:
        if base not in have:
            raise SystemExit(f"style {sid} would be based on {base}, which the reference "
                             f"document does not define.")
    added = 0
    for sid, label, kind, base, unbold in PANDOC_STYLES:
        if sid in have:
            continue
        st = styles.makeelement(qn("w:style"),
                                {qn("w:type"): kind, qn("w:styleId"): sid})
        st.append(st.makeelement(qn("w:name"), {qn("w:val"): label}))
        st.append(st.makeelement(qn("w:basedOn"), {qn("w:val"): base}))
        if unbold:
            rpr = st.makeelement(qn("w:rPr"), {})
            for tag in ("w:b", "w:bCs"):
                rpr.append(rpr.makeelement(qn(tag), {qn("w:val"): "0"}))
            st.append(rpr)
        styles.append(st)
        added += 1
    return added


# Everything V6.tex prints outside sections/. The Nomenclature sits before the body and
# CRediT after it, and both were silently dropped: the builder read only sections/*.tex, and
# a comment here wrongly claimed CRediT came in through frontmatter_body.tex, which does not
# contain it. Elsevier requires CRediT, so this was a desk-check risk.
FRONT_MATTER = ("Nomenclature",)
BACK_MATTER = ("CRediT authorship contribution statement", "Funding", "Acknowledgements",
               "Declaration of competing interest", "Data availability")


def _sections_from_main(names) -> str:
    """Pull named \\section* blocks out of V6.tex, in the order given.

    Taken from V6.tex rather than retyped, so the Word file cannot drift from the PDF.
    A description list is flattened to "term -- definition" lines, because pandoc renders
    elsarticle's \\item[...] glossary as an unlabelled list otherwise.
    """
    src = open(os.path.join(HERE, "V6.tex"), encoding="utf-8").read()
    out, missing = [], []
    for name in names:
        m = re.search(r"\\section\*\{" + re.escape(name)
                      + r"\}(.*?)(?=\\section\*|\\input\{|\\bibliography|\Z)", src, re.S)
        if not m:
            missing.append(name)
            continue
        # Only an unescaped % starts a comment. Stripping \% as well truncated the
        # nomenclature entry for the fixed operating-cost rate mid-formula, which left
        # an unbalanced $ and killed the whole pandoc run.
        body = re.sub(r"(?<!\\)%.*", "", m.group(1))
        if "\\begin{description}" in body:
            items = re.findall(r"\\item\[(.*?)\]\s*(.*?)(?=\\item\[|\\end\{description\})",
                               body, re.S)
            body = "\n\n".join("%s -- %s" % (" ".join(k.split()), " ".join(v.split()))
                                for k, v in items if v.strip())
        else:
            body = " ".join(body.split())
        if body.strip():
            out.append("\\section*{%s}\n%s" % (name, body))
    if missing:
        print("  NOT FOUND in V6.tex:", ", ".join(missing))
    print("  sections carried into Word: %d of %d" % (len(out), len(names)))
    return "\n\n".join(out)


# ── source assembly, reusing build.py's preprocessing so the text is identical ───────
def build_source() -> str:
    sys.path.insert(0, HERE)
    import build as ae_build                                    # noqa: E402
    # The Nomenclature prints before the body in the PDF and belongs there in Word too.
    parts = [_sections_from_main(FRONT_MATTER)]
    for key in ae_build.SECTION_ORDER:
        body = ae_build._read(os.path.join(HERE, "sections", key + ".tex"))
        if body:
            parts.append(ae_build.preprocess(body))
    # Back matter, lifted verbatim from V6.tex. Without this the Word file ended at the
    # bibliography with no competing-interest, funding, acknowledgements or data
    # availability statement, all of which Elsevier requires and the PDF carries. CRediT
    # arrived only because it sits inside frontmatter_body.tex.
    parts.append(_sections_from_main(BACK_MATTER))
    parts.append(r"\section*{References}")
    body = "\n\n".join(parts)

    # Resolve \ref and \eqref against V6.aux. Without this pandoc leaks the raw
    # label, so the Word file read "Eq. [eq:recovery]" where the PDF reads "Eq. (2)".
    # The long paper's build has always done this; this one never did, which is the
    # defect the second author reported.
    body, missing = _resolve_refs(body)
    if missing:
        print("  unresolved labels:", ", ".join(sorted(missing)))

    # Number the captions. pandoc emits a caption with no number at all, which is why
    # the Word file had no "Figure 1:" anywhere while the PDF numbered all seven. The
    # number comes from the same V6.aux the cross-references resolve against, so the
    # two agree by construction.
    body = _number_captions(body)

    # Lift figure note blocks out of the float. pandoc keeps a figure's caption and
    # discards everything else in the environment, so a note set under the figure
    # reached the PDF and vanished from the Word file. Moving it to a paragraph after
    # \end{figure} keeps the two documents saying the same thing.
    body = _lift_figure_notes(body)

    # Mark numbered equations so the Word pass can set the number beside them.
    body = _mark_equations(body)

    src = os.path.join(HERE, "_elsevier_body.tex")
    open(src, "w", encoding="utf-8").write(body)
    return src


AUX_LABEL = re.compile(r"\\newlabel\{([^}]+)\}\{\{([^}]*)\}")


def _aux_labels() -> dict:
    """label -> printed number, read from the compiled V6.aux. Never guessed."""
    path = os.path.join(HERE, "V6.aux")
    if not os.path.exists(path):
        # Returning {} here is what made a rename dangerous. Every \ref went unresolved,
        # no SEQ or REF field was written, and the script still exited 0, so the Word file
        # shipped with dead numbering and the gates of the day all passed. Renaming the
        # submission deletes the old .aux, and so does a fresh clone, so this is the
        # ordinary case rather than an exotic one. Fail here instead, where the cause is
        # obvious, rather than leaving it to be caught downstream.
        raise SystemExit(
            f"{os.path.basename(path)} not found. Compile the LaTeX before building the "
            f"Word twin: pdflatex V5 && bibtex V5 && pdflatex V5 && pdflatex V5")
    out = {}
    for line in open(path, encoding="utf-8", errors="replace"):
        m = AUX_LABEL.search(line)
        if m and m.group(1) not in out:
            out[m.group(1)] = m.group(2)
    return out


def _resolve_refs(text: str):
    labels, missing = _aux_labels(), set()

    def sub(m, paren):
        lab = m.group(1)
        if lab in labels and labels[lab]:
            return "(%s)" % labels[lab] if paren else labels[lab]
        missing.add(lab)
        return m.group(0)

    text = re.sub(r"\\eqref\{([^}]+)\}", lambda m: sub(m, True), text)
    text = re.sub(r"\\ref\{([^}]+)\}", lambda m: sub(m, False), text)
    return text, missing


def _number_captions(text: str) -> str:
    """Prefix each float caption with "Figure N: " or "Table N: " from V6.aux."""
    labels = _aux_labels()

    def sub(m):
        block = m.group(0)
        lab = re.search(r"\\label\{((?:fig|tab):[^}]+)\}", block)
        cap = re.search(r"\\caption\{", block)
        if not (lab and cap):
            return block
        num = labels.get(lab.group(1))
        if not num:
            return block
        kind = "Figure" if lab.group(1).startswith("fig:") else "Table"
        i = cap.end()
        return block[:i] + "%s %s: " % (kind, num) + block[i:]

    return re.sub(r"\\begin\{(figure|table)\}.*?\\end\{\1\}", sub, text, flags=re.S)


# A note block under a figure: \par\vspace{3pt}{\footnotesize\raggedright ... \par}
FIG_NOTE = re.compile(
    r"\\par\\vspace\{[^}]*\}\{\\footnotesize\\raggedright(?P<note>.*?)\\par\}",
    re.S)


def _lift_figure_notes(text: str) -> str:
    """Move each figure's note block to a paragraph just after \\end{figure}.

    pandoc renders a figure as image plus caption and drops the rest of the
    environment, so a note set under the figure is silently lost. Tables are safe,
    because their notes sit after \\end{tabular} in a form pandoc keeps. Only the
    figure form needs lifting, and the note text is carried through unchanged.
    """
    lifted = [0]

    def sub(m):
        block = m.group(0)
        note = FIG_NOTE.search(block)
        if not note:
            return block
        body = " ".join(note.group("note").split())
        block = FIG_NOTE.sub("", block)
        lifted[0] += 1
        return block + "\n\n" + body + "\n"

    out = re.sub(r"\\begin\{figure\}.*?\\end\{figure\}", sub, text, flags=re.S)
    if lifted[0]:
        print("  figure notes lifted out of the float:", lifted[0])
    return out


def _mark_equations(text: str) -> str:
    """Append EQ-MARK-n after each numbered equation, for the Word pass to pick up."""
    labels = _aux_labels()

    def sub(m):
        num = labels.get(m.group(1))
        return m.group(0) if not num else m.group(0) + "\n\nEQ-MARK-" + num + "\n"

    return re.sub(r"\\begin\{equation\}.*?\\label\{([^}]+)\}.*?\\end\{equation\}",
                  sub, text, flags=re.S)


# ── post-processing pandoc cannot do ────────────────────────────────────────────────

M_NS = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def fix_table_layout(d) -> int:
    """Rebuild each data table natively instead of patching what pandoc emitted.

    Pandoc's table markup parses cleanly with python-docx and renders wrongly: the
    paragraph before the table is drawn inside the table's last cell while the table's own
    cells print as loose full-width lines, so each of the two body tables sprawled over
    four near-empty pages. Restyling it, fixing the layout and stamping cell widths all
    left that unchanged, which says the problem is the markup itself rather than any
    property on it.

    So the content is read out and a fresh table is built in its place with python-docx.
    That guarantees well-formed rows, a real style, explicit widths and nothing inherited
    from the converter. The text is copied cell for cell, so nothing is retyped here and
    the table cannot drift from the LaTeX it came from.

    The rules and the column alignment come from the LaTeX. A boxed grid is what Word
    gives by default, and it is not what the PDF shows: booktabs draws no vertical rule at
    all and only three horizontal ones, so the Word twin drew a spreadsheet beside a
    typeset table. Every rule below is one the source asks for.

    Equation containers, which pandoc also emits as tables, are left alone: they are the
    borderless two-column idiom that carries the equation and its number, and they render
    correctly.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
    from docx.table import _Cell
    from docx.text.paragraph import Paragraph

    JUST = {"l": WD_ALIGN_PARAGRAPH.LEFT, "c": WD_ALIGN_PARAGRAPH.CENTER,
            "r": WD_ALIGN_PARAGRAPH.RIGHT}
    specs = _tabular_specs()
    rebuilt = 0
    for t in list(d.tables):
        grid = t._tbl.findall(f"{W_NS}tblGrid/{W_NS}gridCol")
        if len(grid) < 3:
            continue                       # equation container, not a data table
        rows = [_read_row(t, tr, _Cell) for tr in t._tbl.findall(f"{W_NS}tr")]
        ncol = max(sum(span for _, span in row) for row in rows)
        spec = specs[rebuilt]
        nhead, align = spec["head"], spec["align"]

        new = d.add_table(rows=len(rows), cols=ncol)
        new.autofit = False
        # The stub column carries the row labels and the rest carry short values, so an
        # even split wraps every label over three lines. Widths are absolute because the
        # layout is fixed; without both, Word re-fits the table to its content.
        widths = [Inches(2.5)] + [Inches(4.0 / (ncol - 1))] * (ncol - 1)
        _set_grid(new, widths)
        _set_booktabs(new, nhead, spec["rules"], len(rows))
        for r, row in enumerate(rows):
            _keep_row_whole(new.rows[r], header=r < nhead)
            col = 0
            for segs, span in row:
                cell = new.cell(r, col)
                if span > 1:
                    cell = cell.merge(new.cell(r, col + span - 1))
                    for extra in cell.paragraphs[1:]:
                        extra._p.getparent().remove(extra._p)
                cell.width = sum(widths[col:col + span], Inches(0))
                par = cell.paragraphs[0]
                par.paragraph_format.line_spacing = 1.0
                par.paragraph_format.space_after = Pt(0)
                # A spanning row is a sub-heading in these tables and ranges left whatever
                # the column asks for, which is what \multicolumn{5}{l} says in the source.
                par.alignment = JUST[align[col] if span == 1 else "l"]
                # Every row but the last holds the next one to it, which keeps the table
                # and the caption above it on one page.
                par.paragraph_format.keep_with_next = r < len(rows) - 1
                for val, sup, sub, italic in segs:
                    run = par.add_run(val)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(10)
                    # booktabs does not bold a column head and neither does the PDF, so
                    # the head is marked by the rule under it and nothing else. nhead is
                    # still what puts that rule and what repeats the head over a break.
                    run.font.bold = False
                    run.font.italic = italic
                    run.font.superscript = sup
                    run.font.subscript = sub
                col += span
        # The caption sits in its own paragraph above the table and has to travel with it.
        prev = t._tbl.getprevious()
        if prev is not None and prev.tag == f"{W_NS}p":
            Paragraph(prev, None).paragraph_format.keep_with_next = True
        # Put the new table where the old one stood, then drop the old table and the
        # empty paragraph add_table leaves at the end of the body.
        t._tbl.addprevious(new._tbl)
        t._tbl.getparent().remove(t._tbl)
        body = d.element.body
        tail = body.findall(f"{W_NS}p")
        if tail and not "".join(tail[-1].itertext()).strip():
            body.remove(tail[-1])
        rebuilt += 1
    if rebuilt != len(specs):
        raise SystemExit(f"{rebuilt} Word data tables against {len(specs)} multi-column "
                         f"tabulars in the source. The header depths cannot be matched up.")
    return rebuilt


def _read_row(table, tr, cell_cls):
    """One source row as (runs, gridSpan) per cell, merges and run formatting intact.

    Reading row.cells instead repeats a merged cell once per column it covers, which is
    how "Derived independently, not a lever" came out printed five times across the row.
    Reading the cell as one string instead loses the raised note markers and the subscript
    in H2, so the runs are carried out one by one and put back the same way.
    """
    out = []
    for tc in tr.findall(f"{W_NS}tc"):
        if tc.find(f".//{M_NS}oMath") is not None:
            raise SystemExit("a data-table cell holds a maths object, which rebuilding "
                             "the table would drop. Handle it before rebuilding.")
        span = tc.find(f"{W_NS}tcPr/{W_NS}gridSpan")
        n = int(span.get(f"{W_NS}val")) if span is not None else 1
        runs = []
        for par in cell_cls(tc, table).paragraphs:
            for r in par.runs:
                if r.text:
                    runs.append((r.text, bool(r.font.superscript),
                                 bool(r.font.subscript), bool(r.font.italic)))
        out.append((runs, n))
    return out


def bind_figures_to_captions(d) -> int:
    """Hold each figure to the caption under it, so the two cannot land on separate pages.

    Figure 6 sat at the foot of one page with its caption alone at the head of the next
    and a third of the page below the image left blank. Word keeps a paragraph with the
    one after it when asked, and the figure and its caption are two paragraphs.
    """
    bound = 0
    for par in d.paragraphs:
        if par.style.name == "Captioned Figure":
            par.paragraph_format.keep_with_next = True
            bound += 1
    return bound


def _next_paragraph(el):
    """The next body paragraph after el, stepping over bookmarks."""
    from docx.text.paragraph import Paragraph
    nxt = el.getnext()
    while nxt is not None and nxt.tag in (f"{W_NS}bookmarkStart", f"{W_NS}bookmarkEnd"):
        nxt = nxt.getnext()
    return Paragraph(nxt, None) if nxt is not None and nxt.tag == f"{W_NS}p" else None


def fix_note_blocks(d) -> tuple:
    """Set the table-note blocks ragged right and drop the spacer above them.

    A table's notes reach Word as one paragraph of line-broken lines. Justification then
    stretches every line except the last across the full measure, so a one-line note read
    as widely spaced words. The LaTeX sets these blocks \\raggedright, so ranging them
    left matches the PDF as well as fixing the look. The spacer is the empty paragraph
    LaTeX's \\\\[3pt] leaves between the table and its notes, which in Word is a blank
    line the height of a double-spaced body line. Dropping it outright put the notes hard
    against the bottom rule, so the 3 pt the source asks for is given back as space above
    the notes instead.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    ranged = dropped = 0
    for par in list(d.paragraphs):
        breaks = par._p.findall(f".//{W_NS}br")
        if not breaks:
            continue
        if not par.text.strip():
            # A bookmarkEnd for the table's cross-reference target sits between the two,
            # so the table is not always the immediately preceding element.
            prev = par._p.getprevious()
            while prev is not None and prev.tag in (f"{W_NS}bookmarkStart",
                                                    f"{W_NS}bookmarkEnd"):
                prev = prev.getprevious()
            if prev is not None and prev.tag == f"{W_NS}tbl":
                par._p.getparent().remove(par._p)
                dropped += 1
                nxt = _next_paragraph(prev)
                if nxt is not None:
                    nxt.paragraph_format.space_before = Pt(6)
            continue
        if len(breaks) > 1:
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            ranged += 1
    return ranged, dropped


def _keep_row_whole(row, header: bool) -> None:
    """Stop a row splitting over a page, and repeat the column head if the table does.

    Table 1 broke with its first head line stranded at the foot of one page and the rest
    of the table on the next, which reads as two different tables.
    """
    from docx.oxml.ns import qn
    pr = row._tr.get_or_add_trPr()
    pr.append(pr.makeelement(qn("w:cantSplit"), {}))
    if header:
        pr.append(pr.makeelement(qn("w:tblHeader"), {}))


def _set_booktabs(table, nhead: int, rules: set, nrow: int) -> None:
    """Draw the rules the LaTeX draws and no others.

    booktabs sets a heavy rule at the top and foot, a light one under the column head and
    one wherever the source writes \\midrule, and never a vertical rule. Word's default is
    a box around every cell, which turned a typeset table into a spreadsheet. Widths here
    are eighths of a point, so 8 is the 1 pt outer rule and 4 the 0.5 pt inner one.
    """
    from docx.oxml.ns import qn
    pr = table._tbl.tblPr
    for old in pr.findall(qn("w:tblBorders")):
        pr.remove(old)
    borders = pr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(borders.makeelement(qn(f"w:{edge}"), {qn("w:val"): "none",
                                                            qn("w:sz"): "0"}))
    pr.append(borders)
    # Room above and below each rule, which is what \aboverulesep and \belowrulesep buy.
    for old in pr.findall(qn("w:tblCellMar")):
        pr.remove(old)
    mar = pr.makeelement(qn("w:tblCellMar"), {})
    for edge, w in (("top", "40"), ("bottom", "40"), ("left", "0"), ("right", "115")):
        mar.append(mar.makeelement(qn(f"w:{edge}"),
                                   {qn("w:w"): w, qn("w:type"): "dxa"}))
    pr.append(mar)

    def rule(r: int, edge: str, sz: str) -> None:
        for cell in table.rows[r].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            bd = tcPr.find(qn("w:tcBorders"))
            if bd is None:
                bd = tcPr.makeelement(qn("w:tcBorders"), {})
                tcPr.append(bd)
            bd.append(bd.makeelement(qn(f"w:{edge}"),
                                     {qn("w:val"): "single", qn("w:sz"): sz,
                                      qn("w:space"): "0", qn("w:color"): "000000"}))

    rule(0, "top", "8")                                    # \toprule
    rule(nrow - 1, "bottom", "8")                          # \bottomrule
    if nhead < nrow:
        rule(nhead, "top", "4")                            # the rule under the head
    for r in sorted(rules):
        if 0 < r < nrow and r != nhead:
            rule(r, "top", "4")                            # every other \midrule


def _set_grid(table, widths) -> None:
    """Fix the column grid, so the widths asked for are the widths drawn."""
    from docx.oxml.ns import qn
    tbl = table._tbl
    pr = tbl.tblPr
    for old in pr.findall(qn("w:tblLayout")):
        pr.remove(old)
    pr.append(pr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"}))
    grid = tbl.find(f"{W_NS}tblGrid")
    for col, w in zip(grid.findall(f"{W_NS}gridCol"), widths):
        col.set(qn("w:w"), str(int(w.twips)))


def _tabular_specs() -> list:
    """Column alignment, head depth and rule positions per tabular, in source order.

    pandoc keeps none of this. It emits no header marker, so nothing in the converted file
    says where the column head ends; it drops the column spec, so every value came out
    ranged left where the source centres or right-ranges it; and it drops \\midrule, so the
    only way to draw the rules the PDF draws is to read where the source puts them.

    Rows are counted the way the rebuilt table counts them, so "rules" holds the index of
    each row that a \\midrule sits above.
    """
    src = open(os.path.join(HERE, "_elsevier_body.tex"), encoding="utf-8").read()
    out = []
    for spec, body in re.findall(r"\\begin\{tabular\}\{([^}]*)\}(.*?)\\end\{tabular\}",
                                 src, re.S):
        align = re.findall(r"[lcr]", spec)
        if len(align) < 3:
            continue
        head, rules, row = None, set(), 0
        for piece in re.split(r"(\\toprule|\\midrule|\\bottomrule|\\\\)", body):
            if piece == r"\midrule":
                rules.add(row)
                if head is None:
                    head = row
            elif piece == "\\\\":
                row += 1
        out.append({"align": align, "head": max(1, head or 1), "rules": rules})
    return out


def set_page_size_letter(d) -> int:
    """Match the compiled PDF's paper. The twin shipped A4 against a Letter PDF.

    elsarticle is loaded with letterpaper, so V6.pdf is 612 x 792 pt, while pandoc's
    reference document defaults to A4. One submission arrived in two formats on two
    different papers, which is the kind of thing a desk editor notices and no gate here
    measured. Margins are set to one inch, the Word manuscript convention.
    """
    from docx.shared import Inches
    n = 0
    for s in d.sections:
        s.page_width, s.page_height = Inches(8.5), Inches(11)
        s.left_margin = s.right_margin = Inches(1)
        s.top_margin = s.bottom_margin = Inches(1)
        n += 1
    return n


def flatten_empty_base_scripts(d) -> int:
    """Turn an OMML sub/superscript with an empty base into a plain Word script run.

    pandoc renders LaTeX like CO$_2$ as a maths object, because the subscript is maths
    and the CO beside it is not. The object has no base, so pandoc puts a zero-width
    space there. Word draws U+200B as nothing and the result reads correctly, which is
    why every check so far passed: the paragraph text says "tCO2" and the OMML is valid.

    Anything that is not Word draws the zero-width space as a box. Rendering V6.docx
    through LibreOffice put a visible box between the C-O and the 2 in all eighteen of
    them, in "tCO2", "MtCO2", "gCO2", "H2" and "real EUR2024", which is what an editor
    sees in a preview pane.

    A base-less maths object is the wrong markup for a chemical subscript in any case, so
    each one becomes a normal run with vertAlign set. That renders identically in Word
    and correctly everywhere else, and it removes eighteen maths objects from a file that
    only needs maths for its two equations.
    """
    from lxml import etree
    n = 0
    for omath in list(d.element.body.iter(M_NS + "oMath")):
        kids = [k for k in omath if k.tag in (M_NS + "sSub", M_NS + "sSup")]
        if len(kids) != 1 or len(list(omath)) != 1:
            continue
        node = kids[0]
        base = node.find(M_NS + "e")
        script = node.find(M_NS + "sub")
        align = "subscript"
        if script is None:
            script = node.find(M_NS + "sup")
            align = "superscript"
        if base is None or script is None:
            continue
        btxt = "".join(t.text or "" for t in base.iter(M_NS + "t"))
        stxt = "".join(t.text or "" for t in script.iter(M_NS + "t"))
        if btxt.strip("\u200b") or not stxt:
            continue                      # a real base, so leave the maths alone
        run = etree.SubElement(omath.getparent(), W_NS + "r")
        rpr = etree.SubElement(run, W_NS + "rPr")
        etree.SubElement(rpr, W_NS + "vertAlign").set(W_NS + "val", align)
        t = etree.SubElement(run, W_NS + "t")
        t.text = stxt
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        omath.addprevious(run)
        omath.getparent().remove(omath)
        n += 1
    return n



def flatten_letterless_maths(d) -> int:
    """Turn a maths object that holds no letters into plain text.

    pandoc makes maths out of anything between dollars, so "$=0.202$" and "$\\times$"
    arrive as OMML even though nothing about them is mathematical: no variable, no
    structure, just an operator and a number. Word renders them acceptably. LibreOffice
    renders a bare OMML relational as an inverted question mark, so the nomenclature read
    "(\u00bf0.202 tCO2/MWh)" in a preview pane, and a preview pane is what an editor opens
    first.

    Only objects with no letter in them are touched, so a lone italic variable, a
    subscripted symbol and anything carrying a fraction, radical or delimiter keeps its
    maths run and its italics. Both real equations carry letters and are untouched.
    """
    from lxml import etree
    import re as _re
    STRUCT = ("sSub", "sSup", "f", "rad", "nary", "d", "func", "acc", "bar",
              "groupChr", "limLow", "limUpp", "m")
    n = 0
    for omath in list(d.element.body.iter(M_NS + "oMath")):
        if any(e.tag.startswith(M_NS) and e.tag[len(M_NS):] in STRUCT for e in omath.iter()):
            continue
        txt = "".join(t.text or "" for t in omath.iter(M_NS + "t"))
        if not txt or _re.search(r"[A-Za-z]", txt):
            continue
        run = etree.Element(W_NS + "r")
        t = etree.SubElement(run, W_NS + "t")
        t.text = txt
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        omath.addprevious(run)
        omath.getparent().remove(omath)
        n += 1
    return n


def merge_math_runs(d) -> int:
    """Join neighbouring maths runs that share a style, so words print as words.

    pandoc writes \\mathrm{peak} as one run per letter, and both Word and LibreOffice set
    a gap between maths runs. The nomenclature printed "p k , p e a k , c o l d" and the
    methods section printed "C O P" and "S C O P" that way. Merging leaves the maths
    itself untouched: same characters, same style, one run instead of four.
    """
    from lxml import etree

    def style(r):
        pr = r.find(M_NS + "rPr")
        return etree.tostring(pr) if pr is not None else b""

    def plain(r):
        return [c.tag for c in r] in ([M_NS + "rPr", M_NS + "t"], [M_NS + "t"])

    merged = 0
    for omath in list(d.element.body.iter(M_NS + "oMath")):
        for parent in [e for e in omath.iter() if len(e)]:
            i = 0
            while i < len(parent) - 1:
                a, b = parent[i], parent[i + 1]
                # A run holding nothing but space carries no style of its own, so it takes
                # the style beside it. Without this the space between "pk," and "peak"
                # blocked the merge and both kept their maths gap.
                blank = (style(b) == b"" and plain(b)
                         and not (b.find(M_NS + "t").text or "").strip())
                if (a.tag == b.tag == M_NS + "r" and plain(a) and plain(b)
                        and (style(a) == style(b) or blank)):
                    ta, tb = a.find(M_NS + "t"), b.find(M_NS + "t")
                    ta.text = (ta.text or "") + (tb.text or "")
                    if ta.text != ta.text.strip():
                        ta.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    parent.remove(b)
                    merged += 1
                    continue
                i += 1
    return merged


def finalise(fm: dict) -> None:
    import docx
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d = docx.Document(OUT)
    sec = d.sections[0]

    # No margin line numbers, by the author's instruction.

    # Page numbers, centred in the footer.
    p = sec.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    for kind, text in (("begin", None), (None, "PAGE"), ("end", None)):
        if kind:
            fc = OxmlElement("w:fldChar")
            fc.set(qn("w:fldCharType"), kind)
            r._r.append(fc)
        else:
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = text
            r._r.append(it)
    p.style = d.styles["Normal"]
    p.paragraph_format.line_spacing = 1.0

    # Build the title page at the top, in reverse so each insert lands above the last.
    def insert(anchor, text, *, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER,
               italic=False, spacing=LINE_SPACING):
        par = d.add_paragraph()
        run = par.add_run(text)
        run.bold, run.italic = bold, italic
        run.font.size = Pt(BODY_PT)
        run.font.name = BODY_FONT
        par.alignment = align
        par.paragraph_format.line_spacing = spacing
        anchor._p.addprevious(par._p)
        return par

    def insert_image(anchor, path, width_cm):
        par = d.add_paragraph()
        par.add_run().add_picture(path, width=Cm(width_cm))
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        par.paragraph_format.line_spacing = 1.0
        anchor._p.addprevious(par._p)
        return par

    body_start = d.paragraphs[0]
    L, C, J = WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY

    # An author can hold several affiliations, written \author[cornell,jhu]{...}, so
    # the bracket is a tag list rather than a single tag and each one needs its number.
    tagmap = {tag: i for i, (tag, _) in enumerate(fm["affils"], start=1)}
    def marks(tags: str) -> str:
        return ",".join(str(tagmap[x]) for x in
                        (y.strip() for y in tags.split(",")) if x in tagmap)
    names = ", ".join(f"{n}{marks(t_)}" for t_, n in fm["authors"])

    def page_break():
        par = d.add_paragraph()
        par.add_run().add_break(WD_BREAK.PAGE)
        body_start._p.addprevious(par._p)

    # The Word file has to paginate like the PDF, because the two are the same submission
    # and an editor opening either should meet the same thing on the same page. It did
    # not: everything above ran together on one page, in a different order as well
    # (title, abstract, keywords, highlights, graphical abstract), while the PDF gives
    # each of them a page of its own and leads with the graphical abstract.
    #
    # The PDF's order, which this now follows, is: graphical abstract, highlights, title
    # with authors and affiliations, abstract with keywords, then the nomenclature at the
    # head of the body. The nomenclature arrives through FRONT_MATTER, so the last break
    # below is the one that puts it on its own page. check_word_page_structure() in
    # check_submission.py asserts the order and the breaks.
    ga = os.path.join(HERE, "..", "figs", "paper", "graphical_abstract.png")
    if os.path.exists(ga):
        insert(body_start, "Graphical abstract", bold=True, align=L)
        insert(body_start, "", align=L)
        insert_image(body_start, ga, 16.0)
        page_break()

    insert(body_start, "Highlights", bold=True, align=L)
    insert(body_start, "", align=L)
    for h in fm["highlights"]:
        insert(body_start, "\u2022 " + h, align=L)
    page_break()

    insert(body_start, fm["title"], bold=True, align=C)
    insert(body_start, "", align=L)
    insert(body_start, names, align=C)
    for i, (_tag, org) in enumerate(fm["affils"], start=1):
        insert(body_start, f"{i} {org}", align=C, italic=True)
    corr = ", ".join(x for x in (fm["corresponding"], fm["email"]) if x)
    insert(body_start, "", align=L)
    insert(body_start, f"Corresponding author: {corr}", align=C, italic=True)
    page_break()

    insert(body_start, "Abstract", bold=True, align=L)
    insert(body_start, fm["abstract"], align=J)
    insert(body_start, "", align=L)
    insert(body_start, "Keywords: " + "; ".join(fm["keywords"]), align=L)
    page_break()

    # elsarticle renders \paragraph unnumbered, so pandoc's "2.0.0.1" is an
    # artefact of --number-sections reaching a level LaTeX leaves alone. Demote
    # any level-4-or-deeper head to a bold unnumbered lead-in.
    for par in d.paragraphs:
        name = par.style.name
        if name.startswith("Heading") and name[8:].isdigit() and int(name[8:]) >= 4:
            par.style = d.styles["Normal"]
            for run in par.runs:
                run.bold = True
            # Drop the leading number run and the tab that separated it.
            while par.runs and not re.sub(r"^[\d.\s]+", "", par.runs[0].text):
                par.runs[0]._r.getparent().remove(par.runs[0]._r)

    # Every run to the body face and size, so no pandoc default survives anywhere.
    for par in d.paragraphs:
        for run in par.runs:
            run.font.name = BODY_FONT
            if run.font.size is None or run.font.size > Pt(BODY_PT):
                run.font.size = Pt(BODY_PT)
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    # Body face and size inside tables too, as the LaTeX
                    # source shrinks nothing; single-spaced so the grid does
                    # not run past the page.
                    par.paragraph_format.line_spacing = 1.0
                    for run in par.runs:
                        run.font.name = BODY_FONT
                        run.font.size = Pt(BODY_PT)

    # Live Word numbering: SEQ counters on captions, REF fields on cross-references,
    # and numbered equations set in the borderless-table idiom.
    sys.path.insert(0, os.path.dirname(HERE))
    import docx_fields
    st = docx_fields.apply(d)
    print("  fields     : %d equations, %d captions, %d cross-references"
          % (st.get("equations", 0), st.get("captions", 0), st.get("refs", 0)))

    flat = flatten_empty_base_scripts(d)
    lets = flatten_letterless_maths(d)
    runs = merge_math_runs(d)
    print(f"  base-less maths flattened to script runs: {flat}; "
          f"letterless maths flattened to text: {lets}; "
          f"maths runs merged: {runs}")
    tabs = fix_table_layout(d)
    ranged, spacers = fix_note_blocks(d)
    bound = bind_figures_to_captions(d)
    pages = set_page_size_letter(d)
    print(f"  data tables rebuilt natively: {tabs}; note blocks ranged left: {ranged}; "
          f"table spacers dropped: {spacers}; figures bound to captions: {bound}; "
          f"sections set to US Letter: {pages}")
    d.save(OUT)


def main() -> int:
    import pypandoc

    fm = frontmatter()
    make_reference_docx()
    src = build_source()

    args = [
        "--from=latex",
        "--citeproc",
        "--bibliography=%s" % BIB,
        "--resource-path=%s" % os.pathsep.join([HERE, FIGDIR]),
        "--reference-doc=%s" % REF,
        "--number-sections",
    ]
    if os.path.exists(CSL):
        args += ["--csl=%s" % CSL]
    pypandoc.convert_file(src, "docx", outputfile=OUT, extra_args=args)
    finalise(fm)

    print(f"title      : {fm['title'][:70]}...")
    print(f"authors    : {', '.join(n for _, n in fm['authors'])}")
    print(f"abstract   : {len(fm['abstract'].split())} words")
    print(f"highlights : {len(fm['highlights'])}, longest {max(len(h) for h in fm['highlights'])} chars")
    print(f"keywords   : {len(fm['keywords'])}")
    print(f"wrote {OUT} ({os.path.getsize(OUT)} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
