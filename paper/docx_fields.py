"""Live Word numbering for the two Word builds: SEQ captions, REF cross-references,
and equation numbers set the way Word sets them.

WHY THIS EXISTS. pandoc emits a document whose numbers are dead text. Caption numbers
are literal characters, cross-references are literal characters, and numbered equations
lose their number entirely because LaTeX assigns it and pandoc has nowhere to put it.
The result reads correctly until somebody edits it, at which point the numbering is
silently wrong and nothing warns them. A Word-native document instead carries fields:
SEQ for the counters, REF for the pointers, and Word recomputes both.

WHAT EACH PASS DOES.

  equation_numbers()  A staged marker paragraph, EQ-MARK-n, follows every numbered
                      equation. Each is replaced by a borderless one-row table with the
                      equation in the left cell and a right-aligned (n) in the right,
                      which is the idiom Elsevier's own Word templates use and the only
                      one that survives a resize. The number is bookmarked so REF can
                      point at it.

  caption_seq()       A caption reading "Figure 3: ..." keeps its text but the 3 becomes
                      SEQ Figure \\* ARABIC, wrapped in a bookmark.

  ref_fields()        An in-text "Figure 3" becomes REF to that bookmark.

Bookmark names are deterministic (fig_3, tab_2, eq_5) so the three passes agree without
passing state between them.

A field carries both an instruction and a cached result, so the document displays
correctly before Word ever recalculates. Nothing here invents a number: every value comes
from what LaTeX already assigned and pandoc already rendered.
"""
from __future__ import annotations
import re

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

EQ_MARK = re.compile(r"^\s*EQ-MARK-(\d+)\s*$")
CAPTION = re.compile(r"^(Figure|Table)\s+(\d+)([:.])")
KIND_KEY = {"Figure": "fig", "Table": "tab"}


# ── low-level field and bookmark plumbing ───────────────────────────────────────────
def _field_runs(instr: str, cached: str, *, font=None, size=None,
                bold=False, italic=False):
    """The four runs Word needs for a complex field: begin, instruction, result, end.

    The cached result matters. Without it the field shows blank until the reader
    presses F9, which is exactly the "number is missing" complaint this replaces.
    """
    out = []

    def _run(child):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        if font:
            rf = OxmlElement("w:rFonts")
            for a in ("w:ascii", "w:hAnsi", "w:cs"):
                rf.set(qn(a), font)
            rPr.append(rf)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        if size:
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
            rPr.append(sz)
        if len(rPr):
            r.append(rPr)
        r.append(child)
        out.append(r)
        return r

    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin"); _run(b)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = instr; _run(it)
    s = OxmlElement("w:fldChar"); s.set(qn("w:fldCharType"), "separate"); _run(s)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = cached; _run(t)
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); _run(e)
    return out


_BM_ID = [1000]


def _bookmark(name: str):
    _BM_ID[0] += 1
    s = OxmlElement("w:bookmarkStart")
    s.set(qn("w:id"), str(_BM_ID[0])); s.set(qn("w:name"), name)
    e = OxmlElement("w:bookmarkEnd"); e.set(qn("w:id"), str(_BM_ID[0]))
    return s, e


def _run_font(par):
    """Face and size of the paragraph's first real run, so inserted fields match."""
    for r in par.runs:
        if r.text:
            return (r.font.name, r.font.size.pt if r.font.size else None)
    return (None, None)


# ── pass 1: numbered equations ──────────────────────────────────────────────────────
def equation_numbers(doc) -> int:
    """Replace each EQ-MARK-n paragraph by numbering the equation above it."""
    from docx.table import Table

    body = doc.element.body
    marks = [p for p in doc.paragraphs if EQ_MARK.match(p.text)]
    done = 0
    for par in marks:
        n = EQ_MARK.match(par.text).group(1)
        prev = par._p.getprevious()
        while prev is not None and prev.tag != qn("w:p"):
            prev = prev.getprevious()
        if prev is None:
            par._p.getparent().remove(par._p)
            continue

        tbl = OxmlElement("w:tbl")
        tblPr = OxmlElement("w:tblPr")
        w = OxmlElement("w:tblW"); w.set(qn("w:w"), "5000"); w.set(qn("w:type"), "pct")
        tblPr.append(w)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            b = OxmlElement(f"w:{edge}"); b.set(qn("w:val"), "none")
            b.set(qn("w:sz"), "0"); borders.append(b)
        tblPr.append(borders)
        tbl.append(tblPr)
        grid = OxmlElement("w:tblGrid")
        for wd in ("8500", "1000"):
            gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), wd); grid.append(gc)
        tbl.append(grid)

        row = OxmlElement("w:tr")
        left = OxmlElement("w:tc")
        ltcPr = OxmlElement("w:tcPr")
        lw = OxmlElement("w:tcW"); lw.set(qn("w:w"), "8500"); lw.set(qn("w:type"), "dxa")
        ltcPr.append(lw); left.append(ltcPr)
        row.append(left)

        right = OxmlElement("w:tc")
        rtcPr = OxmlElement("w:tcPr")
        rw = OxmlElement("w:tcW"); rw.set(qn("w:w"), "1000"); rw.set(qn("w:type"), "dxa")
        va = OxmlElement("w:vAlign"); va.set(qn("w:val"), "center")
        rtcPr.append(rw); rtcPr.append(va); right.append(rtcPr)
        np = OxmlElement("w:p")
        npPr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "right"); npPr.append(jc)
        np.append(npPr)
        bs, be = _bookmark(f"eq_{n}")
        font, size = _run_font(par)
        # The number prints inside parentheses and the bookmark holds the number alone.
        # Spanning the brackets as well made REF return "(2)", and ref_fields() sets an
        # equation reference in brackets of its own, so "Eq. (2)" came back as
        # "Eq. ((2))" the moment Word or LibreOffice refreshed the field.
        np.append(_text_run("(", font, size))
        np.append(bs)
        for r in _field_runs(f" SEQ Equation \\* ARABIC ", n, font=font, size=size):
            np.append(r)
        np.append(be)
        np.append(_text_run(")", font, size))
        right.append(np)
        row.append(right)
        tbl.append(row)

        prev.addprevious(tbl)
        left.append(prev)                      # move the equation into the left cell
        par._p.getparent().remove(par._p)
        done += 1
    return done


def _text_run(text, font=None, size=None):
    r = OxmlElement("w:r")
    if font or size:
        rPr = OxmlElement("w:rPr")
        if font:
            rf = OxmlElement("w:rFonts")
            for a in ("w:ascii", "w:hAnsi", "w:cs"):
                rf.set(qn(a), font)
            rPr.append(rf)
        if size:
            sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
            rPr.append(sz)
        r.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    r.append(t)
    return r


# ── pass 2: caption numbers as SEQ fields ───────────────────────────────────────────
def caption_seq(doc) -> dict:
    """Turn the literal number in each caption into a SEQ field inside a bookmark.

    Returns {("Figure", "3"): "fig_3"} so ref_fields() knows what exists. Only captions
    whose number matches the running count are converted, because a mismatch means the
    document order and the printed number disagree and a SEQ field would then renumber
    the document wrongly.
    """
    seen = {"Figure": 0, "Table": 0}
    marks = {}
    for par in doc.paragraphs:
        m = CAPTION.match(par.text)
        if not m:
            continue
        kind, num = m.group(1), m.group(2)
        seen[kind] += 1
        if str(seen[kind]) != num:
            continue
        runs = par.runs
        if not runs:
            continue
        # Find the run holding the number and split the caption around it.
        joined, idx, off = "", None, 0
        for i, r in enumerate(runs):
            if idx is None and num in (joined + r.text):
                idx, off = i, (joined + r.text).index(num) - len(joined)
                break
            joined += r.text
        if idx is None:
            continue
        r = runs[idx]
        before, after = r.text[:off], r.text[off + len(num):]
        font = r.font.name
        size = r.font.size.pt if r.font.size else None
        bold = bool(r.font.bold)
        r.text = before
        anchor = r._r
        name = f"{KIND_KEY[kind]}_{num}"
        bs, be = _bookmark(name)
        anchor.addnext(be)
        for el in reversed(_field_runs(f" SEQ {kind} \\* ARABIC ", num,
                                       font=font, size=size, bold=bold)):
            anchor.addnext(el)
        anchor.addnext(bs)
        if after:
            be.addnext(_text_run(after, font, size))
        marks[(kind, num)] = name
    return marks


# ── pass 3: in-text cross-references as REF fields ──────────────────────────────────
def ref_fields(doc, marks: dict) -> int:
    """Rewrite in-text 'Figure 3', 'Table 2' and 'Eq. (5)' as REF fields."""
    # The parentheses are captured, not discarded. Written as `\(?...\)?` for the
    # "Eq. (2)" form, the optional closing bracket also swallowed the one that closes
    # an ordinary parenthetical, so "(Fig. 1)" reached Word as "(Fig. 1." Seven
    # sentences in the submission lost their bracket that way, in the Word file only.
    pat = re.compile(r"\b(Figure|Fig\.|Table|Tab\.|Eq\.|Equation)\s*(\(?)(\d+)(\)?)")
    done = 0
    for par in doc.paragraphs:
        if CAPTION.match(par.text):
            continue                                   # captions carry SEQ, not REF
        for r in list(par.runs):
            text = r.text
            if not text:
                continue
            out, last = [], 0
            for m in pat.finditer(text):
                word, opn, num, cls = m.group(1), m.group(2), m.group(3), m.group(4)
                kind = ("Figure" if word.startswith(("Figure", "Fig"))
                        else "Table" if word.startswith(("Table", "Tab"))
                        else "Equation")
                name = (marks.get((kind, num)) if kind != "Equation"
                        else f"eq_{num}")
                if name is None:
                    continue
                # An equation is always set in parentheses even where the source wrote
                # none; every other reference gets back exactly the brackets it had.
                if kind == "Equation":
                    opn, cls = "(", ")"
                out.append((m.start(), m.end(), word, num, name, kind, opn, cls))
            if not out:
                continue
            font = r.font.name
            size = r.font.size.pt if r.font.size else None
            anchor = r._r
            pieces = []
            for s_, e, word, num, name, kind, opn, cls in out:
                pieces.append(("t", text[last:s_] + word + " " + opn))
                pieces.append(("f", (name, num, font, size)))
                pieces.append(("t", cls))
                last = e
            pieces.append(("t", text[last:]))
            r.text = ""
            for what, payload in reversed(pieces):
                if what == "t":
                    if payload:
                        anchor.addnext(_text_run(payload, font, size))
                else:
                    name, num, f_, s_ = payload
                    for el in reversed(_field_runs(f" REF {name} \\h ", num,
                                                   font=f_, size=s_)):
                        anchor.addnext(el)
                    done += 1
    return done


def apply(doc, *, equations=True, captions=True, refs=True) -> dict:
    """Run the three passes and report what each one changed."""
    stats = {}
    if equations:
        stats["equations"] = equation_numbers(doc)
    marks = caption_seq(doc) if captions else {}
    stats["captions"] = len(marks)
    if refs:
        stats["refs"] = ref_fields(doc, marks)
    return stats
