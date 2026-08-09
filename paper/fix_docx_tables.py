"""Post-process the pandoc-generated docx: full text-width tables, black link text, and
strip the source-repository sentence from the Word output.

Tables. Pandoc emits tables at tblW=auto with equal narrow columns, so text-heavy columns
clip. Simple (rectangular) tables are set to Word's "AutoFit to Window" (tblW=100% pct plus
an autofit layout): the table fills the text area and columns size to their content. Tables
that contain vertically merged cells (\\multirow in the source, e.g. the cost-optimal
robustness table) are NOT autofit: autofit recomputes column widths and collapses the merged
label column. Those are instead pinned to a fixed layout at an explicit dxa width equal to the
text area, with the source column proportions preserved, so the merged cells render cleanly.

Links. hyperref is set to hidelinks so the PDF link text is black; for parity the Word build
forces every hyperlink run to black here (belt-and-braces over the Hyperlink character style,
which make_ref_docx.py also sets to black).

Repository URL. The "full model code and data are available at <github link>" sentence is
removed from the Word output (the link is stripped, not just recoloured).

Run: python fix_docx_tables.py in.docx out.docx
"""
import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

W = qn("w:w")
TYPE = qn("w:type")
VAL = qn("w:val")


def _child(parent, tag):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _set_tblW(pr, type_, w):
    tblW = pr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        pr.insert(0, tblW)
    tblW.set(TYPE, type_)
    tblW.set(W, str(w))


def _set_layout(pr, type_):
    layout = pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        pr.append(layout)
    layout.set(TYPE, type_)


def text_area_twips(doc):
    # python-docx Length.__int__ is EMUs; gridCol/tblW dxa values are twips, so use .twips
    s = doc.sections[0]
    return s.page_width.twips - s.left_margin.twips - s.right_margin.twips


def fix_tables(doc):
    """Full-width simple tables (autofit); fixed-width vMerge tables (honour proportions)."""
    area = text_area_twips(doc)
    n_simple = n_fixed = 0
    for t in doc.tables:
        tbl = t._tbl
        pr = tbl.tblPr
        has_vmerge = bool(tbl.findall(".//" + qn("w:vMerge")))
        if has_vmerge:
            # (a) fixed layout at the text-area width; rescale gridCols to keep proportions.
            t.autofit = False
            grid = tbl.find(qn("w:tblGrid"))
            cols = grid.findall(qn("w:gridCol")) if grid is not None else []
            widths = [int(c.get(W)) for c in cols if c.get(W)]
            total = sum(widths) or area
            acc = 0
            for i, c in enumerate(cols):
                if i == len(cols) - 1:
                    neww = area - acc                       # last column absorbs rounding
                else:
                    neww = round(int(c.get(W)) * area / total)
                    acc += neww
                c.set(W, str(neww))
            # pin every cell's tcW to its column width so fixed layout is honoured
            for row in tbl.findall(qn("w:tr")):
                tcs = row.findall(qn("w:tc"))
                cw = [int(c.get(W)) for c in cols]
                ci = 0
                for tc in tcs:
                    tcPr = _child(tc, "w:tcPr")
                    span = tcPr.find(qn("w:gridSpan"))
                    nspan = int(span.get(VAL)) if span is not None and span.get(VAL) else 1
                    wsum = sum(cw[ci:ci + nspan]) if ci < len(cw) else area
                    tcW = _child(tcPr, "w:tcW")
                    tcW.set(TYPE, "dxa")
                    tcW.set(W, str(wsum))
                    ci += nspan
            _set_tblW(pr, "dxa", area)
            _set_layout(pr, "fixed")
            n_fixed += 1
        else:
            # simple rectangular table: AutoFit to Window (fills text area, content-based cols)
            t.autofit = True
            _set_tblW(pr, "pct", 5000)                      # 5000 fiftieths = 100%
            n_simple += 1
    return n_simple, n_fixed


def links_black(doc):
    """Force every hyperlink run to black text (parity with the PDF's hidelinks)."""
    body = doc.element.body
    n = 0
    for hl in body.iter(qn("w:hyperlink")):
        for r in hl.findall(qn("w:r")):
            rPr = _child(r, "w:rPr")
            color = rPr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rPr.append(color)
            color.set(VAL, "000000")
            n += 1
    return n


def strip_repo_sentence(doc):
    """Remove 'The full model code and data are available at <github link>.' from the body.

    Anchored on the run text and the GitHub hyperlink so no other paragraph is touched.
    """
    LEAD = "The full model code and data are available at"
    removed = 0
    for p in doc.paragraphs:
        children = list(p._p)
        # is there a hyperlink to github in this paragraph?
        gh = None
        for el in children:
            if el.tag == qn("w:hyperlink"):
                for r in el.findall(qn("w:r")):
                    txt = "".join(t.text or "" for t in r.findall(qn("w:t")))
                    if "github.com/alajwadha" in txt:
                        gh = el
                        break
            if gh is not None:
                break
        if gh is None:
            continue
        # find the lead-in run; remove from there to the end of the sentence (incl. the link
        # and the trailing period run that follows the hyperlink)
        lead_idx = None
        for i, el in enumerate(children):
            if el.tag == qn("w:r"):
                txt = "".join(t.text or "" for t in el.findall(qn("w:t")))
                if txt.strip().startswith(LEAD):
                    lead_idx = i
                    break
        if lead_idx is None:
            continue
        gh_idx = children.index(gh)
        end_idx = gh_idx
        # absorb a single trailing period run right after the hyperlink
        if gh_idx + 1 < len(children) and children[gh_idx + 1].tag == qn("w:r"):
            nxt = children[gh_idx + 1]
            ntxt = "".join(t.text or "" for t in nxt.findall(qn("w:t")))
            if ntxt.strip() in (".", ". "):
                end_idx = gh_idx + 1
        for el in children[lead_idx:end_idx + 1]:
            p._p.remove(el)
        # the run just before the lead-in is "...accommodation tables [163]." (already ends the
        # prior sentence with its own period); drop a dangling separator space if one remains
        prev = list(p._p)
        if lead_idx - 1 >= 0 and lead_idx - 1 < len(prev) and prev[lead_idx - 1].tag == qn("w:r"):
            ptxt = "".join(t.text or "" for t in prev[lead_idx - 1].findall(qn("w:t")))
            if ptxt == " ":
                p._p.remove(prev[lead_idx - 1])
        removed += 1
    return removed


def main():
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else sys.argv[1]
    doc = Document(src)
    n_simple, n_fixed = fix_tables(doc)
    n_links = links_black(doc)
    n_repo = strip_repo_sentence(doc)
    doc.save(dst)
    print(f"{len(doc.tables)} tables: {n_simple} full-width autofit, {n_fixed} fixed-width "
          f"(vMerge); {n_links} hyperlink runs set black; repo sentence removed x{n_repo}")


if __name__ == "__main__":
    main()
