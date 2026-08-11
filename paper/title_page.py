"""Finish the Word document: one clean native title page, and A/B/C appendix lettering.

Title page: replace pandoc's split title block with a single centred native title page (title,
authors, affiliations, date), followed by a page break so the abstract starts on page 2.

Appendix lettering: pandoc's --number-sections numbers every heading sequentially in arabic, so
the appendices come out as 8, 9, 10 instead of the PDF's A, B, C. This re-letters the appendix
Heading 1s (and their A.1/A.2 children) to match the LaTeX \\appendix output. The boundary is
found from the first appendix section title, so the offset adapts if body sections are added.

Run: python title_page.py Paper_v20.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn

INK = RGBColor(0x1A, 0x1A, 0x1A)

# title of the first appendix section (the \appendix boundary). Used to locate where arabic
# numbering should switch to A/B/C lettering. If this is ever not found, lettering is skipped
# (headings stay numbered) and a notice is printed, so the build never breaks.
FIRST_APPENDIX_TITLE = "Supplementary Methodology and Detailed Results"


def build_title_page(doc):
    # 1. drop pandoc's title/author/date paragraphs at the top
    for p in list(doc.paragraphs[:8]):
        if p.style and p.style.name in ("Title", "Author", "Date"):
            p._p.getparent().remove(p._p)

    made = []

    def add(text, size, bold=False, italic=False, after=6, before=0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_after, pf.space_before, pf.line_spacing = Pt(after), Pt(before), 1.15
        r = p.add_run(text)
        f = r.font
        f.name, f.size, f.bold, f.italic, f.color.rgb = "Cambria", Pt(size), bold, italic, INK
        made.append(p._p)
        return p

    add("Hydrogen in Residential Heating:", 20, bold=True, after=2, before=84)
    add("A Merit-Order Assessment for the EU, the United Kingdom and", 16, bold=True, after=2)
    add("Switzerland under Net-Zero by 2050", 16, bold=True, after=30)
    add("Abdurahman Alsulaiman¹˒*      Ali Alajwad²˒³", 13, after=16)
    add("¹ Laboratory of Environmental and Urban Economics (LEURE), EPFL, Switzerland",
        11, italic=True, after=2)
    add("² RENEW Lab, Johns Hopkins University, Baltimore, MD, USA", 11, italic=True, after=2)
    add("³ Cornell University, Ithaca, NY, USA", 11, italic=True, after=6)
    add("* Corresponding author.", 11, italic=True, after=18)
    last = add("Draft v20  ·  June 2026", 11, after=0)
    last.add_run().add_break(WD_BREAK.PAGE)

    body = doc.element.body
    anchor = body[0]
    for el in made:
        body.remove(el)
    for el in made:                       # forward order, each inserted just before the anchor
        anchor.addprevious(el)


def _first_text_run(p):
    """The first w:r in the paragraph that carries a w:t (where pandoc puts the number)."""
    for r in p._p.findall(qn("w:r")):
        if r.findall(qn("w:t")):
            return r
    return None


def _run_text(r):
    return "".join(t.text or "" for t in r.findall(qn("w:t")))


def _set_run_text(r, value):
    ts = r.findall(qn("w:t"))
    ts[0].text = value
    for extra in ts[1:]:
        extra.text = ""


def reletter_appendices(doc):
    """Convert appendix headings from arabic (8, 8.1) to letters (A, A.1) to match the PDF."""
    heads = [p for p in doc.paragraphs
             if p.style and p.style.name in ("Heading 1", "Heading 2")]
    # locate the first appendix Heading 1 by its title
    start = None
    for i, p in enumerate(heads):
        if p.style.name == "Heading 1" and FIRST_APPENDIX_TITLE in p.text:
            start = i
            break
    if start is None:
        print("title_page: first appendix title not found; left headings numbered")
        return 0
    # the arabic number on that first appendix Heading 1 fixes the offset (e.g. 8 -> A)
    m0 = re.match(r"\s*(\d+)", heads[start].text)
    if not m0:
        print("title_page: first appendix heading has no number; left headings numbered")
        return 0
    base = int(m0.group(1))                       # arabic number of appendix A
    changed = 0
    for p in heads[start:]:
        r = _first_text_run(p)
        if r is None:
            continue
        txt = _run_text(r)
        m = re.match(r"^(\d+)(\.\d+)?$", txt.strip())
        if not m:
            continue
        top = int(m.group(1))
        letter = chr(ord("A") + top - base)       # 8->A, 9->B, ...
        if not ("A" <= letter <= "Z"):
            continue
        new = letter + (m.group(2) or "")          # 'A' or 'A.1'
        # preserve any leading/trailing whitespace the run had
        lead = txt[:len(txt) - len(txt.lstrip())]
        trail = txt[len(txt.rstrip()):]
        _set_run_text(r, f"{lead}{new}{trail}")
        changed += 1
    print(f"title_page: re-lettered {changed} appendix headings (A/B/C to match the PDF)")
    return changed


def main():
    path = sys.argv[1]
    doc = Document(path)
    build_title_page(doc)
    reletter_appendices(doc)
    doc.save(path)
    print("prepended a clean native title page")


if __name__ == "__main__":
    main()
